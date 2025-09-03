"""
模块名称: test_doc.py

该模块的目标：
    编写bp测试的word报告
  
作者: ych
修改历史:
    1. 2025/9/3 - 创建文件
"""
import sys
from docx.shared import Inches, Pt, RGBColor
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn  # 用于处理中文字体


def custom_heading(doc, text, level=1):
    """
    自定义样式的标题

    参数:
        doc: Document 对象
        text: 标题文本
        level: 标题级别 (1-9)
    """
    # 添加标题
    heading = doc.add_heading(level=level)
    # 添加run并设置文本
    run = heading.add_run(text)
    # 设置字体为宋体
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置字号
    if level == 1:
        run.font.size = Pt(21)  # 一级标题21号
    elif level == 2:
        run.font.size = Pt(18)  # 二级标题18号
    else:
        run.font.size = Pt(16)  # 其他标题16号
    # 设置加粗
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # RGB(0,0,0) 表示黑色
    return heading


def add_unified_paragraph(doc, text):
    """
    添加统一格式的段落：宋体、11号、黑色、不加粗
    参数:
        doc: Document 对象
        text: 段落文本
    """
    # 添加段落
    paragraph = doc.add_paragraph()
    # 添加run并设置文本
    run = paragraph.add_run(text)
    # 设置字体为宋体
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置字号为11号
    run.font.size = Pt(11)
    # 设置黑色
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 确保不加粗
    run.font.bold = False
    return paragraph


def test_doc():
    # 获取当前日期
    today = datetime.now()
    # 提取月份和日期
    month = today.month  # 直接获取月份（1-12）
    day = today.day  # 直接获取日期（1-31）
    formatted_date = f"{month}-{day}"
    last_part = ip.rsplit(".", 1)[-1]  # 从IP地址中提取最后一部分
    # 创建新文档
    doc = Document()
    # 添加标题
    custom_heading(doc, f'{formatted_date}-bp测试', 1)
    custom_heading(doc, '正常-网桥直通模式切换', level=1)
    custom_heading(doc, '正常模式', level=3)
    add_unified_paragraph(doc, '仪表盘：')
    doc.add_picture('./img/1.png', width=Inches(6.0))
    # 保存文档
    doc.save(f'./doc_files/{formatted_date}-bp测试({last_part}).docx')


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("用法: python test_doc.py <输入IP>")
        print("示例: python test_doc.py 10.20.192.106")
        sys.exit(1)

    ip = sys.argv[1]
    test_doc()
